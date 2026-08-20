#!/usr/bin/env python3
"""Generate the public one-page, tagged resume PDF from shared JSON content."""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "content" / "resume.public.json"
OUTPUT_PATH = ROOT / "public" / "Jesadakorn-Kirtnu-Resume.pdf"

FONT_NAME = "Arial"
NAVY = RGBColor.from_string("273B51")
INK = RGBColor.from_string("18212B")
ACCENT = RGBColor.from_string("B86F43")

PDF_EXPORT_OPTIONS = {
    "PDFUACompliance": {"type": "boolean", "value": "true"},
    "UseTaggedPDF": {"type": "boolean", "value": "true"},
    "ExportBookmarks": {"type": "boolean", "value": "true"},
    "DisplayPDFDocumentTitle": {"type": "boolean", "value": "true"},
    "EnableTextAccessForAccessibilityTools": {"type": "boolean", "value": "true"},
}
PDF_EXPORT_FILTER = (
    "pdf:writer_pdf_Export:" + json.dumps(PDF_EXPORT_OPTIONS, separators=(",", ":"))
)


def load_data() -> dict:
    with DATA_PATH.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def set_ooxml_font_size(run_properties, size: float) -> None:
    half_points = f"{size * 2:g}"
    for tag in ("w:sz", "w:szCs"):
        size_element = run_properties.find(qn(tag))
        if size_element is None:
            size_element = OxmlElement(tag)
            run_properties.append(size_element)
        size_element.set(qn("w:val"), half_points)


def set_style_font(style, *, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.font.italic = False
    style.font.underline = False

    run_properties = style.element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attribute}"), FONT_NAME)

    set_ooxml_font_size(run_properties, size)
    for inherited_tag in ("w:spacing", "w:kern"):
        inherited_element = run_properties.find(qn(inherited_tag))
        if inherited_element is not None:
            run_properties.remove(inherited_element)

    language = run_properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run_properties.append(language)
    language.set(qn("w:val"), "en-US")


def format_run(
    run,
    *,
    size: float = 9.25,
    color: RGBColor = INK,
    bold: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline

    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    set_ooxml_font_size(run_properties, size)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a descriptive, externally linked run to a paragraph."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = paragraph.add_run(text)
    format_run(run, color=ACCENT, underline=True)
    run_element = run._element
    paragraph._element.remove(run_element)
    hyperlink.append(run_element)
    paragraph._element.append(hyperlink)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal, size=9.25, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = Pt(10.6)
    normal.paragraph_format.widow_control = False

    title = styles["Title"]
    set_style_font(title, size=20, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing = Pt(21)
    title.paragraph_format.keep_with_next = True
    title_properties = title.element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=10, color=ACCENT, bold=True)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle.paragraph_format.line_spacing = Pt(11.5)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_properties = subtitle.element.get_or_add_pPr()
    subtitle_numbering = subtitle_properties.find(qn("w:numPr"))
    if subtitle_numbering is not None:
        subtitle_properties.remove(subtitle_numbering)

    heading_one = styles["Heading 1"]
    set_style_font(heading_one, size=10.5, color=NAVY, bold=True)
    heading_one.paragraph_format.space_before = Pt(4)
    heading_one.paragraph_format.space_after = Pt(1)
    heading_one.paragraph_format.line_spacing = Pt(11.5)
    heading_one.paragraph_format.keep_with_next = True
    heading_one.paragraph_format.keep_together = True

    heading_two = styles["Heading 2"]
    set_style_font(heading_two, size=9.5, color=NAVY, bold=True)
    heading_two.paragraph_format.space_before = Pt(2)
    heading_two.paragraph_format.space_after = Pt(0)
    heading_two.paragraph_format.line_spacing = Pt(10.5)
    heading_two.paragraph_format.keep_with_next = True
    heading_two.paragraph_format.keep_together = True

    list_bullet = styles["List Bullet"]
    set_style_font(list_bullet, size=9.25, color=INK)
    list_bullet.paragraph_format.left_indent = Inches(0.2)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.12)
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(0)
    list_bullet.paragraph_format.line_spacing = Pt(10.4)
    list_bullet.paragraph_format.widow_control = False


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_section_heading(document: Document, text: str) -> None:
    document.add_paragraph(text.upper(), style="Heading 1")


def build_resume_docx(data: dict, output_path: Path) -> None:
    document = Document()
    configure_document(document)

    properties = document.core_properties
    properties.title = f"{data['identity']['name']} - Public Resume"
    properties.subject = "Technology operations, software, and automation professional resume"
    properties.author = data["identity"]["name"]
    properties.keywords = "technology operations, software, automation, resume"
    properties.comments = "ATS-first, one-page public resume generated from shared portfolio data"

    name = document.add_paragraph(style="Title")
    name.add_run(data["identity"]["name"])

    headline = document.add_paragraph(style="Subtitle")
    headline.add_run(data["identity"]["headline"])

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(2)
    contact.add_run(f"{data['identity']['location']}  |  ")
    descriptive_link_text = {
        "Portfolio": "Portfolio: jesadakorn.com",
        "LinkedIn": f"LinkedIn: {data['identity']['name']}",
        "GitHub": "GitHub: fluke-jesadakorn",
    }
    for index, link in enumerate(data["links"]):
        if index:
            contact.add_run("  |  ")
        add_hyperlink(
            contact,
            descriptive_link_text.get(link["label"], link["label"]),
            link["url"],
        )

    add_section_heading(document, "Professional Summary")
    summary = document.add_paragraph(data["summary"])
    summary.paragraph_format.space_after = Pt(1)

    add_section_heading(document, "Core Capabilities")
    for capability in data["expertise"]:
        add_bullet(document, capability)

    add_section_heading(document, "Experience")
    for experience in data["experience"]:
        role = document.add_paragraph(style="Heading 2")
        role_run = role.add_run(experience["role"])
        format_run(role_run, size=9.5, color=NAVY, bold=True)
        organization_run = role.add_run(f" | {experience['organization']}")
        format_run(organization_run, size=9.25, color=INK, bold=False)
        period_run = role.add_run(f" | {experience['period']}")
        format_run(period_run, size=9.25, color=ACCENT, bold=True)
        for bullet in experience["bullets"]:
            add_bullet(document, bullet)

    add_section_heading(document, "Selected Systems")
    for system in data["selectedSystems"]:
        add_bullet(document, system)

    add_section_heading(document, "Education & Qualifications")
    education = data["education"]
    education_paragraph = document.add_paragraph()
    education_paragraph.paragraph_format.space_after = Pt(0)
    degree = education_paragraph.add_run(education["degree"])
    format_run(degree, bold=True, color=NAVY)
    education_paragraph.add_run(
        f" | {education['institution']} | {education['period']}"
    )
    add_bullet(document, f"Languages: {' and '.join(data['languages'])}")
    for qualification in data["additional"]:
        add_bullet(document, qualification)

    add_section_heading(document, "Technology")
    technology = document.add_paragraph(", ".join(data["technologies"]))
    technology.paragraph_format.space_after = Pt(0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def find_libreoffice() -> str:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if executable:
        return executable

    macos_path = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if macos_path.exists():
        return str(macos_path)

    raise RuntimeError(
        "LibreOffice is required to generate the resume PDF. Install it and make "
        "soffice or libreoffice available on PATH."
    )


def export_pdf(docx_path: Path, output_path: Path, work_dir: Path) -> None:
    libreoffice = find_libreoffice()
    profile_dir = work_dir / "libreoffice-profile"
    home_dir = work_dir / "libreoffice-home"
    profile_dir.mkdir()
    home_dir.mkdir()

    environment = os.environ.copy()
    environment["HOME"] = str(home_dir)
    environment["TMPDIR"] = str(work_dir)

    result = subprocess.run(
        [
            libreoffice,
            "--headless",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--convert-to",
            PDF_EXPORT_FILTER,
            "--outdir",
            str(work_dir),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        env=environment,
    )
    generated_pdf = work_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not generated_pdf.exists():
        details = "\n".join(part for part in (result.stdout, result.stderr) if part.strip())
        raise RuntimeError(f"LibreOffice PDF export failed.\n{details}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(generated_pdf, output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--docx-output",
        type=Path,
        help="Optionally retain a copy of the semantic DOCX for accessibility auditing.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = load_data()
    with tempfile.TemporaryDirectory(prefix="jesadakorn-resume-") as temporary_directory:
        work_dir = Path(temporary_directory)
        docx_path = work_dir / "Jesadakorn-Kirtnu-Resume.docx"
        build_resume_docx(data, docx_path)

        if args.docx_output:
            retained_docx = args.docx_output.resolve()
            retained_docx.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(docx_path, retained_docx)
            print(f"Generated accessibility-audit DOCX at {retained_docx}")

        export_pdf(docx_path, OUTPUT_PATH, work_dir)

    print(f"Generated {OUTPUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
